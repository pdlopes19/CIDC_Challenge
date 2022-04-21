from tkinter import *
from tkinter import filedialog
import os
from PyQt5 import QtCore, QtGui, QtWidgets
import PyQt5
import pandas as pd
import time
import sys
from datetime import datetime
from docx import Document

from cidcchallenge import Info
from cidcchallenge import Spreadsheet
from cidcchallenge import Operators
from cidcchallenge import Services
from cidcchallenge import Employees

from Start import Ui_button_start
from Operators import Ui_Operators
from Services import Ui_Services
from Employees import Ui_Employees
from Main_Menu import Ui_Main_Window
from Menu import Ui_Second_Window

from Graphics_Menu import Ui_Graphics_Menu
from Graphics_Employees import Ui_Graphics_Employees
from Graphics_Histograms import Ui_Histograms
from Graphics_Operators import Ui_Graphics_Operators
from Graphics_Services import Ui_Graphics_Services
from Finish import Ui_Finish


class GUI(object):

    window = None

    archive_name = None

    information = None # Info's class object
    spreadsheet = None # Spreadsheet's class object

    data_frame = None # Pandas Data Frame, contains all the data provenient of the spreadsheet
    # Important Columns = 'Account', 'Activity Type', 'Duration', 'Assigned to'

    data_frame_wo = None # Pandas Data Frame, contains all the data provenient of the spreadsheet excluding outliers
    # Columns = 'Operator', 'Employee', 'Duration [s]', 'Task', 'Activity'

    timestamps = None # Timestamps list with outlier
    timestamp_designation = None # Timestamps designation list
    # [0] Employees [1] Timestamps [2] Task [3] Activity [4] Operator

    std_deviation = None # Standard deviation with outliers
    mean = None # Mean with outliers

    std_deviation_wo = None # Standard deviation without outliers
    mean_wo = None # Mean without outliers

    operators_list = None
    services_list = None
    employees_list = None

    operators_number_of_services = None # Dictionary - keys on operators_list
    operators_number_of_services_wo = None
    operators_number_of_unique_services = None # Dictionary - keys = operator + service
    operators_number_of_unique_services_wo = None
    operators_mean = None

    services_number_of_unique_services = None # Dictionary - keys on services_list
    services_number_of_unique_services_wo = None
    services_mean = None

    employees_above_below_mean = None # Dictionary - keys on employees_list
    employees_above_below_mean_wo = None
    employees_mean = None


    def __init__(self, _archive_name):

        self.archive_name = _archive_name
        self.get_values_from_cidcchallenge()

    def get_values_from_cidcchallenge(self):

        self.data_frame = pd.read_csv(self.archive_name + '.csv')

        # Instance of Info class, will give us the basic informations
        self.information = Info(self.archive_name)
        self.information.show_basic_info()

        # Calculus of the number informations englobing operators and services
        self.information.calc_number_info_operators()
        self.information.calc_number_info_services()

        # Instance of spreadsheet, will transform the spreadsheet for better execution of operations
        self.spreadsheet = Spreadsheet(self.archive_name)
        self.timestamps = self.spreadsheet.timestamp_conversion()
        self.timestamp_designation = self.spreadsheet.timestamp_designation()

        # Calculus of std deviation and mean, considering outliers
        self.std_deviation = self.spreadsheet.calc_std_deviation_with_outlier()
        self.mean = self.spreadsheet.calc_mean_with_outlier()

        # Receive the outlier exclusion pattern, on the case 3 times the standard deviation
        # outlier_pattern = int(input('Outlier exclusion pattern (times standard deviation): '))
        outlier_pattern = 3

        # Exclusion of outliers
        self.data_frame_wo = self.spreadsheet.exclude_outliers(outlier_pattern)

        # Calculus of std deviation and mean, desconsidering outliers
        self.std_deviation_wo = self.spreadsheet.calc_std_deviation_without_outlier()
        self.mean_wo = self.spreadsheet.calc_mean_without_outlier()

        # Calculus of the number informations englobing employees
        self.timestamp_designation = self.spreadsheet.timestamp_designation()
        self.information.calc_number_info_employees_with_outlier(self.mean, self.timestamp_designation)
        self.information.calc_number_info_employees_without_outlier(self.mean_wo,
                                                                    self.data_frame_wo)

        self.operators_number_of_services = self.information.operators_obj.count_number_of_services()

        self.operators_number_of_services_wo = self.information.operators_obj.count_number_of_services_without_outlier(
            self.data_frame_wo)

        self.operators_number_of_unique_services = self.information.operators_obj.count_number_of_unique_services_per_operator()

        self.operators_number_of_unique_services_wo = self.information.operators_obj.count_number_of_unique_services_per_operator_without_outlier(
            self.data_frame_wo)

        self.operators_mean = self.information.operators_obj.calc_operators_mean_duration(self.data_frame_wo)

        self.services_number_of_unique_services_wo = self.information.services_obj.count_number_of_unique_services_without_outlier(
            self.data_frame_wo)

        self.services_number_of_unique_services = self.information.services_obj.count_number_of_unique_services()

        self.services_mean = self.information.services_obj.calc_services_mean_duration(self.data_frame_wo)

        self.employees_above_below_mean = self.information.employees_obj.calc_above_below_mean_employees_with_outlier(
            self.mean, self.timestamp_designation)

        self.employees_above_below_mean_wo = self.information.employees_obj.calc_above_below_mean_employees_without_outlier(
            self.mean_wo, self.data_frame_wo)

        self.employees_mean = self.information.employees_obj.calc_mean_duration_per_employee(self.data_frame_wo)
        self.services_mean = self.information.services_obj.calc_services_mean_duration(self.data_frame_wo)

        self.operators_list = self.information.operators_obj.show_operators()
        self.services_list = self.information.services_obj.show_services()
        self.employees_list = self.information.employees_obj.show_employees()

    def show_values_from_cidcchallenge(self):

        print(self.data_frame)
        print(self.data_frame_wo)

        print(self.timestamps)
        print(self.timestamp_designation)

        print(self.std_deviation)
        print(self.mean)

        print(self.std_deviation_wo)
        print(self.mean_wo)

        print(self.operators_list)
        print(self.services_list)
        print(self.employees_list)

        print(self.operators_number_of_services)
        print(self.operators_number_of_services_wo)
        print(self.operators_number_of_unique_services)
        print(self.operators_number_of_unique_services_wo)
        print(self.operators_mean)

        print(self.services_number_of_unique_services)
        print(self.services_number_of_unique_services_wo)
        print(self.services_mean)

        print(self.employees_above_below_mean)
        print(self.employees_above_below_mean_wo)

class Operations(object):

    archive_name = None
    gui = None

    def __init__(self):

        self.app = QtWidgets.QApplication(sys.argv)
        self.button_start = QtWidgets.QMainWindow()
        self.ui = Ui_button_start()
        self.ui.setupUi(self.button_start)
        self.button_start.show()
        self.ui.button_startProject.clicked.connect(self.menu)
        sys.exit(self.app.exec_())

    def menu(self):

        self.window = QtWidgets.QDialog()
        self.ui = Ui_Main_Window()
        self.ui.setupUi(self.window)
        self.window.show()
        self.button_start.destroy()
        self.ui.button_ok_mainMenu.clicked.connect(self.get_archive)

    def get_archive(self):

        self.ui.search.setAcceptRichText(False)
        text = self.ui.search.toPlainText()

        self.archive_name = text
        print(self.archive_name)

        self.window.close()

        self.gui = GUI(self.archive_name)

        self.window = QtWidgets.QDialog()
        self.ui = Ui_Second_Window()
        self.ui.setupUi(self.window)
        self.window.show()

        self.ui.button_operators.clicked.connect(self.operators_menu)
        self.ui.button_services.clicked.connect(self.services_menu)
        self.ui.button_employees.clicked.connect(self.employees_menu)
        self.ui.button_graphics.clicked.connect(self.graphics_menu)
        self.ui.button_close.clicked.connect(self.finish)

    def operators_menu(self):

        self.window.close()

        self.window = QtWidgets.QDialog()
        self.ui = Ui_Operators()
        self.ui.setupUi(self.window)
        self.window.show()

        text_with_outliers = ''
        text_without_outliers = ''

        text_without_outliers += '[Number of services / operator]\n\n'
        text_with_outliers += '[Number of services / operator]\n\n'

        for i in self.gui.operators_list:

            text_with_outliers += '{}: {}\n'.format(i, self.gui.operators_number_of_services[i])
            text_without_outliers += '{}: {}\n'.format(i, self.gui.operators_number_of_services_wo[i])

        text_with_outliers += '\n'
        text_without_outliers += '\n'

        text_with_outliers += '[Number of unique services / operator]\n\n'
        text_without_outliers += '[Number of unique services / operator]\n\n'

        for i in self.gui.operators_list:

            for j in self.gui.services_list:

                key = str(i) + ' + ' + str(j)

                text_with_outliers += '{}: {}\n'.format(key, self.gui.operators_number_of_unique_services[key])
                text_without_outliers += '{}: {}\n'.format(key, self.gui.operators_number_of_unique_services_wo[key])

        text_with_outliers += '\n'
        text_without_outliers += '\n'

        text_without_outliers += '[Services mean time / operator]\n\n'

        for i in self.gui.operators_list:

            timestamp = time.strftime('%H:%M:%S', time.gmtime(self.gui.operators_mean[i]))
            text_without_outliers += '{}: {}\n'.format(i, timestamp)

        self.ui.text_operators.setText(text_with_outliers)
        self.ui.text_woOperators.setText(text_without_outliers)

        self.ui.button_backOperators.clicked.connect(self.back_menu)

    def back_menu(self):

        self.window.close()

        self.window = QtWidgets.QDialog()
        self.ui = Ui_Second_Window()
        self.ui.setupUi(self.window)
        self.window.show()

        self.ui.button_operators.clicked.connect(self.operators_menu)
        self.ui.button_services.clicked.connect(self.services_menu)
        self.ui.button_employees.clicked.connect(self.employees_menu)
        self.ui.button_graphics.clicked.connect(self.graphics_menu)
        self.ui.button_close.clicked.connect(self.finish)


    def services_menu(self):

        self.window.close()
        self.window = QtWidgets.QDialog()
        self.ui = Ui_Services()
        self.ui.setupUi(self.window)
        self.window.show()

        text_with_outliers = ''
        text_without_outliers = ''

        text_with_outliers += '[Number of unique services]\n\n'
        text_without_outliers += '[Number of unique services]\n\n'

        for i in self.gui.services_list:

            text_with_outliers += '{}: {}\n'.format(i, self.gui.services_number_of_unique_services[i])
            text_without_outliers += '{}: {}\n'.format(i, self.gui.services_number_of_unique_services_wo[i])

        text_with_outliers += '\n'
        text_without_outliers += '\n'

        text_without_outliers += '[Mean duration of each service]\n\n'


        for i in self.gui.services_list:

            try:

                timestamp = time.strftime('%H:%M:%S', time.gmtime(self.gui.services_mean[i]))
            except:

                timestamp = '00:00:00'

            text_without_outliers += '{}: {}\n'.format(i, timestamp)

        self.ui.text_services.setText(text_with_outliers)
        self.ui.text_woServices.setText(text_without_outliers)

        self.ui.button_services.clicked.connect(self.back_menu)

    def employees_menu(self):

        self.window.close()

        self.window = QtWidgets.QDialog()
        self.ui = Ui_Employees()
        self.ui.setupUi(self.window)
        self.window.show()

        above = {}
        above_wo = {}
        below = {}
        below_wo = {}

        employees_column = list(self.gui.employees_above_below_mean['Employee'])
        employees_column_wo = list(self.gui.employees_above_below_mean_wo['Employee'])
        above_mean_column = list(self.gui.employees_above_below_mean['Above mean?'])
        above_mean_column_wo = list(self.gui.employees_above_below_mean_wo['Above mean?'])

        for i in self.gui.employees_list:

            above[i] = 0
            above_wo[i] = 0
            below[i] = 0
            below_wo[i] = 0

        for i in range(0, len(employees_column)):

            if above_mean_column[i] == 'Yes':

                above[employees_column[i]] += 1

            else:

                below[employees_column[i]] += 1

        for i in range(0, len(employees_column_wo)):

            if above_mean_column_wo[i] == 'Yes':

                above_wo[employees_column_wo[i]] += 1

            else:

                below_wo[employees_column_wo[i]] += 1

        text = ''
        text_wo = ''

        text += '[Above mean]\n\n'
        text_wo += '[Above mean]\n\n'

        for i in self.gui.employees_list:

            text += '{}: above mean {} times\n'.format(i, above[i])
            text_wo += '{}: above mean {} times\n'.format(i, above_wo[i])

        text += '\n[Below mean]\n\n'
        text_wo += '\n[Below mean]\n\n'

        for i in self.gui.employees_list:

            text += '{}: below mean {} times\n'.format(i, below[i])
            text_wo += '{}: below mean {} times\n'.format(i, below_wo[i])

        text_wo += '\n[Employees mean]\n\n'

        for i in self.gui.employees_list:

            timestamp = time.strftime('%H:%M:%S', time.gmtime(self.gui.employees_mean[i]))
            text_wo += '{}: {}\n'.format(i, timestamp)

        self.ui.text_employees.setText(text)
        self.ui.text_woEmployees.setText(text_wo)

        self.ui.button_employees.clicked.connect(self.back_menu)

    def graphics_menu(self):

        self.window.close()

        self.window = QtWidgets.QDialog()
        self.ui = Ui_Graphics_Menu()
        self.ui.setupUi(self.window)
        self.window.show()

        self.ui.button_operatorsG.clicked.connect(self.graphics_menu_operators)
        self.ui.button_servicesG.clicked.connect(self.graphics_menu_services)
        self.ui.button_employeesG.clicked.connect(self.graphics_menu_employees)
        self.ui.button_histograms.clicked.connect(self.graphics_histograms)
        self.ui.button_backG.clicked.connect(self.back_menu)


        # self.gui.information.show_graphic_with_outliers(self.gui.timestamps, self.gui.mean, self.gui.std_deviation)
        # self.gui.information.show_graphic_without_outliers(self.gui.data_frame_wo['Duration [s]'], self.gui.mean_wo, self.gui.std_deviation_wo)
        # self.gui.information.show_graphic_operators(self.gui.operators_number_of_services, 'considering outliers')
        # self.gui.information.show_graphic_operators(self.gui.operators_number_of_services_wo, 'desconsidering outliers')
        # self.gui.information.show_graphic_operators_unique_services(self.gui.operators_number_of_unique_services, 'considering outliers')
        # self.gui.information.show_graphic_operators_unique_services(self.gui.operators_number_of_unique_services_wo, 'desconsidering outliers')
        # self.gui.information.show_graphic_operators_mean(self.gui.operators_mean, 'desconsidering outliers')
        # self.gui.information.show_graphic_services_unique(self.gui.services_number_of_unique_services, 'considering outliers')
        # self.gui.information.show_graphic_services_unique(self.gui.services_number_of_unique_services_wo, 'desconsidering outliers')
        # self.gui.information.show_graphic_employees_mean(self.gui.employees_mean)

    def graphics_menu_operators(self):

        self.window.close()
        self.window = QtWidgets.QDialog()
        self.ui = Ui_Graphics_Operators()
        self.ui.setupUi(self.window)
        self.window.show()

        self.ui.button_CO.clicked.connect(self.operators1)
        self.ui.button_DO.clicked.connect(self.operators2)
        self.ui.button_uniqueSer_CO.clicked.connect(self.operators3)
        self.ui.button_uniqueSer_DO.clicked.connect(self.operators4)
        self.ui.button_meanOperators.clicked.connect(self.operators5)
        self.ui.button_backOperators.clicked.connect(self.back_menu)

    def operators1(self):

        self.gui.information.show_graphic_operators(self.gui.operators_number_of_services, 'considering outliers')

    def operators2(self):

        self.gui.information.show_graphic_operators(self.gui.operators_number_of_services_wo, 'desconsidering outliers')

    def operators3(self):

        self.gui.information.show_graphic_operators_unique_services(self.gui.operators_number_of_unique_services,
                                                                    'considering outliers')

    def operators4(self):

        self.gui.information.show_graphic_operators_unique_services(self.gui.operators_number_of_unique_services_wo,
                                                                    'desconsidering outliers')

    def operators5(self):

        self.gui.information.show_graphic_operators_mean(self.gui.operators_mean, 'desconsidering outliers')

    def graphics_menu_services(self):

        self.window.close()
        self.window = QtWidgets.QDialog()
        self.ui = Ui_Graphics_Services()
        self.ui.setupUi(self.window)
        self.window.show()

        self.ui.button_services_CO.clicked.connect(self.services1)
        self.ui.button_services_DO.clicked.connect(self.services2)
        self.ui.button_backOperators.clicked.connect(self.back_menu)

    def services1(self):

        self.gui.information.show_graphic_services_unique(self.gui.services_number_of_unique_services,
                                                          'considering outliers')

    def services2(self):

        self.gui.information.show_graphic_services_unique(self.gui.services_number_of_unique_services_wo,
                                                          'desconsidering outliers')

    def graphics_menu_employees(self):

        self.window.close()
        self.window = QtWidgets.QDialog()
        self.ui = Ui_Graphics_Employees()
        self.ui.setupUi(self.window)
        self.window.show()

        self.ui.button_meanEmployees.clicked.connect(self.employees1)
        self.ui.button_backEmployees.clicked.connect(self.back_menu)

    def employees1(self):

        self.gui.information.show_graphic_employees_mean(self.gui.employees_mean, self.gui.mean_wo, self.gui.std_deviation_wo)

    def graphics_histograms(self):

        self.window.close()
        self.window = QtWidgets.QDialog()
        self.ui = Ui_Histograms()
        self.ui.setupUi(self.window)
        self.window.show()

        self.ui.button_histograms_CO.clicked.connect(self.histograms1)
        self.ui.button_histograms_DO.clicked.connect(self.histograms2)
        self.ui.button_backEmployees.clicked.connect(self.back_menu)


    def histograms1(self):

        self.gui.information.show_graphic_with_outliers(self.gui.timestamps, self.gui.mean, self.gui.std_deviation)

    def histograms2(self):

        self.gui.information.show_graphic_without_outliers(self.gui.data_frame_wo['Duration [s]'], self.gui.mean_wo,
                                                           self.gui.std_deviation_wo)

    def finish(self):

        self.window.close()
        self.window = QtWidgets.QDialog()
        self.ui = Ui_Finish()
        self.ui.setupUi(self.window)
        self.window.show()

        self.ui.button_closeProject.clicked.connect(self.close)

    def close(self):

        self.window.close()

        login = 'Pedro Lopes de Oliveira'
        date = datetime.now()
        date = date.strftime('%d/%m/%y - %H:%M')
        doc = Document()
        doc.add_heading('CIDC Challenge - Building a report through a spread sheet', 0)
        doc.add_paragraph('This software was developed by Pedro Lopes de Oliveira and Yara Caroline Tavares Mendes, Undergraduated Telecommunications Engineers.\n\nFor contact:\nPedro - pedro.lopes@get.inatel.br\nYara - yaracaroline@get.inatel.br')
        doc.add_paragraph('Generated by: {}'.format(login) + '\nDate: {}'.format(date))
        doc.add_paragraph('For better quality informations, the outliers were excluded considering the stipulated outlier pattern.')

        doc.add_heading('Operators', 1)

        doc.add_heading('Number of services for each operator', 2)

        doc.add_paragraph(
            'On this subsection of Operators are the information about the number of services executed for each operator. ')

        table = doc.add_table(len(self.gui.operators_number_of_services_wo),
                              len(self.gui.operators_number_of_services_wo))
        table.style = 'Light Grid Accent 1'
        heading_cells = table.rows[0].cells
        heading_cells[0].text = 'Service'
        heading_cells[1].text = 'Number of services'

        for i in self.gui.operators_list:
            cells = table.add_row().cells
            cells[0].text = i
            cells[1].text = str(self.gui.operators_number_of_services_wo[i])

        doc.add_heading('Number of unique services for each operator', 2)

        doc.add_paragraph('On this subsection of Operators are the information about the number of unique services (all services with all operators) executed for each operator. ')

        table = doc.add_table(2,2)
        table.style = 'Light Grid Accent 1'
        heading_cells = table.rows[0].cells
        heading_cells[0].text = 'Service'
        heading_cells[1].text = 'Number of services'

        for i in self.gui.operators_list:
            for j in self.gui.services_list:

                key = i + ' + ' + j
                cells = table.add_row().cells
                cells[0].text = str(key)
                cells[1].text = str(self.gui.operators_number_of_unique_services_wo[key])

        doc.add_heading('Mean duration for each operator', 2)

        doc.add_paragraph(
            'On this subsection of Operators are the information about the mean time used for each operator. ')

        table = doc.add_table(2,2)
        table.style = 'Light Grid Accent 1'
        heading_cells = table.rows[0].cells
        heading_cells[0].text = 'Operator'
        heading_cells[1].text = 'Mean duration'

        for i in self.gui.operators_list:

            cells = table.add_row().cells
            timestamp = time.strftime('%H:%M:%S', time.gmtime(self.gui.operators_mean[i]))
            cells[0].text = i
            cells[1].text = str(timestamp)

        doc.add_heading('Services', 1)

        doc.add_heading('Number of unique services', 2)

        doc.add_paragraph(
            'On this subsection of Services are the information about the number of services executed, desconsidering the operators. ')

        table = doc.add_table(2, 2)
        table.style = 'Light Grid Accent 1'
        heading_cells = table.rows[0].cells
        heading_cells[0].text = 'Service'
        heading_cells[1].text = 'Number of services'

        for i in self.gui.services_list:
            cells = table.add_row().cells
            cells[0].text = i
            cells[1].text = str(self.gui.services_number_of_unique_services_wo[i])

        doc.add_heading('Mean time for each service', 2)

        doc.add_paragraph(
            'On this subsection of Services are the information about the mean time used for each service. ')

        table = doc.add_table(2, 2)
        table.style = 'Light Grid Accent 1'
        heading_cells = table.rows[0].cells
        heading_cells[0].text = 'Service'
        heading_cells[1].text = 'Mean duration'

        for i in self.gui.services_list:
            cells = table.add_row().cells

            try:

                timestamp = time.strftime('%H:%M:%S', time.gmtime(self.gui.services_mean[i]))

            except:

                timestamp = '00:00:00'

            cells[0].text = i
            cells[1].text = str(timestamp)


        doc.add_heading('Employees', 1)

        doc.add_heading('Above mean employees', 2)

        doc.add_paragraph(
            'On this subsection of Employees are the information about the number of times an employee was above and below the mean duration. ')

        table = doc.add_table(2, 2)
        table.style = 'Light Grid Accent 1'
        heading_cells = table.rows[0].cells
        heading_cells[0].text = 'Employee'
        heading_cells[1].text = 'Number of times above mean'

        employee = list(self.gui.employees_above_below_mean_wo['Employee'])
        above = list(self.gui.employees_above_below_mean_wo['Above mean?'])

        info = {}

        for i in self.gui.employees_list:

            info[i] = 0


        for i in range(len(employee)):

            if above[i] == 'Yes':

                info[employee[i]] += 1


        for i in self.gui.employees_list:
            cells = table.add_row().cells
            cells[0].text = i
            cells[1].text = str(info[i])

        doc.add_paragraph('')

        table = doc.add_table(2, 2)
        table.style = 'Light Grid Accent 1'
        heading_cells = table.rows[0].cells
        heading_cells[0].text = 'Employee'
        heading_cells[1].text = 'Number of times below mean'

        for i in self.gui.employees_list:

            info[i] = 0

        for i in range(len(employee)):

            if above[i] == 'No':

                info[employee[i]] += 1


        for i in self.gui.employees_list:
            cells = table.add_row().cells
            cells[0].text = i
            cells[1].text = str(info[i])

        doc.add_heading('Mean time of each employee', 2)

        doc.add_paragraph(
            'On this subsection of Employees are the information about the mean time of each employee to execute activities.')

        table = doc.add_table(2, 2)
        table.style = 'Light Grid Accent 1'
        heading_cells = table.rows[0].cells
        heading_cells[0].text = 'Employee'
        heading_cells[1].text = 'Mean duration (overall)'

        for i in self.gui.employees_list:
            cells = table.add_row().cells

            try:

                timestamp = time.strftime('%H:%M:%S', time.gmtime(self.gui.employees_mean[i]))

            except:

                timestamp = 0

            cells[0].text = i
            cells[1].text = str(timestamp)

        doc.add_paragraph('After these information catched, you should observe that Jorge was above mean on the overall, what is considered a bad profile.')
        doc.add_paragraph('In this report was shown basic informations that the spreadsheet contains, for more information you should check the below folders (generated with the software):\n- data_frames\n- graphics\n- text_archives ')
        doc.add_paragraph('Thank you.')
        doc.add_paragraph('Instituto Nacional de Telecomunicações - ' + date)

        doc.add_page_break()
        doc.save('report.docx')

        exit(0)


if __name__ == '__main__':

    operations = Operations()

