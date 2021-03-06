from docx import Document
from datetime import datetime

login = 'Pedro Lopes de Oliveira'
date = datetime.now()
date = date.strftime('%d/%m/%y - %H:%M')

services_list = list()
services_list.append('1800 MOCN')
services_list.append('Antenna Swap')
services = {'1800 MOCN': 2, 'Antenna Swap': 5}

doc = Document()
doc.add_heading('CIDC Challenge - Building a report through a spread sheet', 0)
doc.add_paragraph('Generated by: {}'.format(login))
doc.add_paragraph('Date: {}'.format(date))

doc.add_heading('Operators',1)
doc.add_heading('Number of services',2)

table = doc.add_table(len(services), len(services))
table.style = 'Light Grid Accent 1'
heading_cells = table.rows[0].cells
heading_cells[0].text = 'Service'
heading_cells[1].text = 'Number of services'


for i in services_list:
    cells = table.add_row().cells
    cells[0].text = i
    cells[1].text = str(services[i])

doc.add_heading('Services',1)

doc.add_heading('Employees',1)

doc.add_page_break()
doc.save('report.docx')